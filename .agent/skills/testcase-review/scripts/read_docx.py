#!/usr/bin/env python3
"""
Đọc tài liệu Word (.docx), trích xuất nội dung văn bản
"""

import json
import sys
import argparse
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Lỗi: Chưa cài đặt python-docx. Hãy chạy: pip install python-docx")
    sys.exit(1)


def read_docx(file_path: str) -> dict:
    """
    Đọc tài liệu Word

    Tham số:
        file_path: Đường dẫn tệp docx

    Trả về:
        Từ điển chứa nội dung và cấu trúc của tài liệu
    """
    doc = Document(file_path)

    content = []
    current_section = None
    sections = {}

    # Trích xuất các đoạn văn (paragraphs)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Xác định xem có phải là tiêu đề (Heading) không
        if para.style.name.startswith('Heading'):
            current_section = text
            if current_section not in sections:
                sections[current_section] = []
        elif current_section:
            sections[current_section].append(text)
        else:
            # Văn bản không có tiêu đề
            if 'Chính văn' not in sections:
                sections['Chính văn'] = []
            sections['Chính văn'].append(text)

        content.append({
            'type': 'heading' if para.style.name.startswith('Heading') else 'paragraph',
            'style': para.style.name,
            'text': text
        })

    # Trích xuất bảng (tables)
    tables = []
    for idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append({
            'index': idx,
            'data': table_data
        })

    # Gộp tất cả văn bản
    all_text = '\n'.join([item['text'] for item in content])

    return {
        'sections': sections,
        'content': content,
        'tables': tables,
        'full_text': all_text
    }


def main():
    parser = argparse.ArgumentParser(description='Đọc tài liệu Word')
    parser.add_argument('file_path', help='Đường dẫn tệp docx')
    parser.add_argument('-o', '--output', help='Đường dẫn xuất tệp JSON')
    parser.add_argument('--text-only', action='store_true', help='Chỉ xuất văn bản thuần túy')

    args = parser.parse_args()

    if not Path(args.file_path).exists():
        print(f"Lỗi: Không tìm thấy tệp: {args.file_path}")
        sys.exit(1)

    result = read_docx(args.file_path)

    if args.text_only:
        print(result['full_text'])
    elif args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"Đã ghi kết quả đầu ra vào: {args.output}")
    else:
        print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
